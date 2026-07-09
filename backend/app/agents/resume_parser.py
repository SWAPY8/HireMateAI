import json
import re
import io
import logging
from pypdf import PdfReader
import docx
from app.core.ai import query_gemini

logger = logging.getLogger("app.resume_parser")

class ResumeParser:
    @staticmethod
    def extract_text(filename: str, content_bytes: bytes) -> str:
        text = ""
        ext = filename.split(".")[-1].lower()
        
        try:
            if ext == "pdf":
                pdf_file = io.BytesIO(content_bytes)
                reader = PdfReader(pdf_file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            elif ext in ["docx", "doc"]:
                docx_file = io.BytesIO(content_bytes)
                doc = docx.Document(docx_file)
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text]
                        if row_text:
                            text += " | ".join(row_text) + "\n"
            else:
                # Text format or general decoding
                text = content_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}", exc_info=True)
            text = ""

        # If extraction is empty or garbage, fall back to decoding
        if not text.strip():
            try:
                text = content_bytes.decode("utf-8", errors="ignore")
            except:
                text = f"Resume filename: {filename}"
                
        return text

    @staticmethod
    def parse(filename: str, file_content: bytes = b"", api_key: str = None) -> dict:
        text = ResumeParser.extract_text(filename, file_content)
        
        # Highly detailed prompt for extraction and parsing
        prompt = f"""
        You are a premium AI Resume Parser and ATS Analyzer. Extract and analyze every detail of the resume text.
        
        Resume Text:
        \"\"\"{text}\"\"\"
        
        Provide the response as a JSON object with the following fields:
        1. "name": Full name of candidate.
        2. "email": Email address.
        3. "phone": Phone number.
        4. "location": City, State, Country location.
        5. "dob": Date of birth if present.
        6. "linkedin": LinkedIn URL.
        7. "github": GitHub URL.
        8. "portfolio": Portfolio website URL.
        9. "summary": Engaging professional summary.
        10. "skills": Comma-separated list of all technical skills.
        11. "programming_languages": List of programming languages (e.g. ["Python", "JS"]).
        12. "frameworks": List of frameworks (e.g. ["React", "FastAPI"]).
        13. "libraries": List of libraries (e.g. ["Redux", "Pydantic"]).
        14. "tools": List of tools/software.
        15. "databases": List of databases.
        16. "cloud_platforms": List of cloud platforms.
        17. "projects": List of projects. Each project should be a JSON object with "title", "description", and "tech_used".
        18. "internships": List of internship experiences. Each should have "company", "role", "duration", and "description".
        19. "experience": List of professional work experiences. Each should have "company", "role", "duration", and "description".
        20. "education": List of education credentials. Each should have "institution", "degree", "field_of_study", and "duration".
        21. "certifications": List of certifications.
        22. "achievements": List of achievements.
        23. "publications": List of publications if present.
        24. "languages_known": List of languages known.
        25. "soft_skills": List of soft skills.
        26. "preferred_role": Target job role recommendation.
        27. "expected_salary": Estimated expected monthly salary in ₹ (e.g. "₹1,20,000").
        28. "preferred_location": Target/preferred work location.
        29. "work_preference": Remote, Hybrid, or On-site.
        30. "notice_period": Notice period estimate.
        
        ATS Analysis & Insights:
        31. "ats_score": Overall ATS match compatibility score as an integer (0 to 100).
        32. "resume_quality_score": General structural and visual quality score (0 to 100).
        33. "keyword_match": Assessment statement explaining how well keywords match typical target JDs.
        34. "missing_skills": Complementary technical skills missing but recommended.
        35. "skill_gap_analysis": Analysis of candidate's skills vs industry needs.
        36. "strongest_sections": List of strong sections (e.g. ["Projects", "Skills"]).
        37. "weakest_sections": List of weak sections requiring development.
        38. "formatting_issues": List of structural or layout formatting issues.
        39. "grammar_suggestions": List of specific grammatical or language suggestions.
        40. "quantifiable_achievement_suggestions": List of ideas to add measurable metrics.
        41. "recruiter_impression": Recruiter's general review of the resume.
        42. "salary_estimate": Recommended monthly package estimate in Indian Rupees (₹) (e.g. "₹1,10,000").
        43. "interview_readiness_score": Interview readiness estimation (0 to 100).
        44. "improvement_roadmap": Structured, numbered step-by-step roadmap to enhance the resume.

        Ensure you return valid JSON. Do not prefix with markdown formatting like ```json.
        """
        
        logger.info(f"Parsing resume {filename} with Gemini API...")
        raw_res = query_gemini(prompt, json_mode=True, api_key=api_key)
        
        fallback = {
            "name": filename.split(".")[0].replace("_", " ").replace("-", " ").title(),
            "email": "candidate@example.com",
            "phone": "Not specified",
            "location": "Not specified",
            "dob": "Not specified",
            "linkedin": "",
            "github": "",
            "portfolio": "",
            "summary": "Experienced professional.",
            "skills": "Python, React, SQL",
            "programming_languages": ["Python", "JavaScript"],
            "frameworks": ["React", "FastAPI"],
            "libraries": [],
            "tools": ["Git"],
            "databases": ["PostgreSQL"],
            "cloud_platforms": ["AWS"],
            "projects": [],
            "internships": [],
            "experience": [],
            "education": [],
            "certifications": [],
            "achievements": [],
            "publications": [],
            "languages_known": ["English"],
            "soft_skills": ["Communication"],
            "preferred_role": "Full Stack Developer",
            "expected_salary": "₹80,000",
            "preferred_location": "Remote",
            "work_preference": "Remote",
            "notice_period": "Immediate",
            "ats_score": 65,
            "resume_quality_score": 60,
            "keyword_match": "General matching level.",
            "missing_skills": ["Docker"],
            "skill_gap_analysis": "Improve containerization skills.",
            "strongest_sections": ["Experience"],
            "weakest_sections": ["Projects"],
            "formatting_issues": ["No page numbers"],
            "grammar_suggestions": [],
            "quantifiable_achievement_suggestions": ["Specify team size."],
            "recruiter_impression": "Strong core developer skillset.",
            "salary_estimate": "₹90,000",
            "interview_readiness_score": 70,
            "improvement_roadmap": "1. Add deployment projects.\n2. Quantify achievements."
        }
        
        if not raw_res:
            logger.error("Gemini API returned empty response for resume parse.")
            raise ValueError("AI Resume Parsing failed. Service is temporarily unavailable.")
            
        try:
            parsed = json.loads(raw_res)
            parsed["raw_text"] = text
            # Coerce scores
            for field in ["ats_score", "resume_quality_score", "interview_readiness_score"]:
                try:
                    parsed[field] = int(parsed.get(field, 65))
                except:
                    parsed[field] = 65
            return parsed
        except Exception as e:
            logger.error(f"Failed to parse JSON response from Gemini API: {e}\nRaw Response: {raw_res}")
            raise ValueError("AI Resume Parsing failed due to invalid response structure.")
